import os
import jwt as pyjwt
from datetime import datetime, timedelta
import openai
import re
import numpy as np
import requests
from flask import Flask, request, jsonify, send_from_directory
from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from sklearn.metrics.pairwise import cosine_similarity
from context_manager import generate_reply_template
from docx import Document as DocxDocument
from langchain.schema import Document
import hashlib
import pytz

# from SQLite_database import (
from Postgres import (
    init_db,
    register_user,
    get_user_by_password,
    get_username,
    store_chat_embedding,
    get_chat_history,
    get_recent_chat_embeddings,
    store_document_embedding,
    get_document_embeddings,
    store_file_metadata,
    get_user_files,
    get_existing_files,
    delete_file,
    get_used_storage
)
 
# Load OpenAI API Key
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")
 
# Initialize Flask App
app = Flask(__name__)
# socketio = SocketIO(app, cors_allowed_origins="*",
#                    logger=False,
#                    engineio_logger=False,
#                    async_mode='threading')
SECRET_KEY = "your_secret_key"
 
# Initialize Embedding Model
embeddings_model = OpenAIEmbeddings(model="text-embedding-ada-002")
# Initialize AI Model
chat_model = ChatOpenAI(model="gpt-4o-mini")
 
# List of pronouns to track
PRONOUNS = {"he", "she", "his", "her", "him", "hers", "they", "them", "their", "theirs",
            "it", "its", "here", "there", "this", "that", "these", "those"}
USER_PRONOUNS = {"me", "my", "mine", "myself"}
 
# Resolve Pronouns
def resolve_pronouns(question, chat_history):
    try:
        if not chat_history:
            return chat_history, question, "true"
 
        last_message = chat_history[0]
        words = set(re.findall(r'\b\w+\b', question.lower()))
 
        if words & PRONOUNS:
            modified_question = chat_model.invoke(
                f"I will give you a message and a question containing pronouns. Rewrite the question by replacing pronouns with the correct entity from the message.\n\nMessage: \"{last_message}\"\nQuestion: \"{question}\"\nRewritten Question:").content
            return [last_message], modified_question, "true"
 
        if words & USER_PRONOUNS:
            chat_response = chat_model.invoke(
                f"Sentence:{question},RETURN ONLY BOOLEAN ,Determine if a given sentence is asking about the user themselves, If the sentence is about the user , return false. Otherwise, return true.").content
            # return chat_history, question, chat_response
 
        return chat_history, question, chat_response
    except Exception as e:
        print(f"Error resolving pronouns: {e}")
        return chat_history, question, "true"
 
# JWT Authentication
def generate_token(user_id):
    try:
        payload = {
            "user_id": user_id,
            "exp": datetime.utcnow() + timedelta(hours=24)
        }
        return pyjwt.encode(payload, SECRET_KEY, algorithm="HS256")
    except Exception as e:
        print(f"Error generating token: {e}")
        return None
 
def verify_token(token):
    try:
        decoded = pyjwt.decode(token, SECRET_KEY, algorithms=["HS256"])
        return decoded["user_id"]
    except pyjwt.ExpiredSignatureError:
        return None
    except pyjwt.InvalidTokenError:
        return None
 
# File Operations
def allowed_file(filename):
    ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS
 
# def extract_text(file_path, file_extension):
#     try:
#         if file_extension == "pdf":
#             loader = PyMuPDFLoader(file_path)
#             return loader.load()
#         else:
#             extracted_text = textract.process(file_path).decode("utf-8")
#             return [Document(page_content=extracted_text)]
#     except Exception as e:
#         print(f"Error extracting text from file: {e}")
#         return []

def extract_text(file_path, file_extension):
    try:
        if file_extension == "pdf":
            loader = PyMuPDFLoader(file_path)
            return loader.load()

        elif file_extension == "docx":
            doc = DocxDocument(file_path)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
            return [Document(page_content=extracted_text)]

        elif file_extension == "txt":
            with open(file_path, "r", encoding="utf-8") as file:
                extracted_text = file.read()
            return [Document(page_content=extracted_text)]

    except Exception as e:
        print(f"Error extracting text from file: {e}")
        return []

 
def compute_file_hash(file_path):
    try:
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            while chunk := f.read(8192):
                hasher.update(chunk)
        return hasher.hexdigest()
    except Exception as e:
        print(f"Error computing file hash: {e}")
        return None
 
MAX_USER_STORAGE = 30 * 1024 * 1024  # 30MB limit
 
def retrieve_similar_context(user_id, newchat_id, query, top_k=1):
    def get_top_matches(data):
        similarities = []
        query_embedding = embeddings_model.embed_query(query)
        query_vector = np.array(query_embedding, dtype=np.float32).reshape(1, -1)
        for text, embedding_blob in data:
            stored_embedding = np.frombuffer(embedding_blob, dtype=np.float32).reshape(1, -1)
            similarity = cosine_similarity(query_vector, stored_embedding)[0][0]
            similarities.append((similarity, text))
        similarities.sort(reverse=True, key=lambda x: x[0])
        return [text for _, text in similarities[:top_k]]
 
    chat_data = get_recent_chat_embeddings(user_id, newchat_id)
    chat_history = [row[0] for row in chat_data]
    relevant_chat_history, modified_query, chat_response = resolve_pronouns(query, chat_history)
    doc_context = ''
    meta_context = ''
    top_match = ''
   
    if (chat_response.lower() == "true"):
        doc_data = get_document_embeddings(user_id)
        top_match = get_top_matches(doc_data)
        doc_context = "\n".join([row[0] for row in doc_data])
        meta_data = get_user_files(user_id)
        meta_context = [{"name": row["file_name"], "doc_preview_link": row["file_path"]} for row in meta_data]
 
    def get_system_location():
        try:
            response = requests.get("http://ip-api.com/json/")
            data = response.json()
            location_info = {
                "Country": data.get("country"),
                "Region": data.get("regionName"),
                "City": data.get("city"),
                "Latitude": data.get("lat"),
                "Longitude": data.get("lon"),
                "ISP": data.get("isp"),
                "IP Address": data.get("query")
            }
            return location_info
        except Exception as e:
            return {"error": str(e)}
 
    location_data = get_system_location()
    context_template = generate_reply_template(
        previous_conversations=relevant_chat_history,
        other_information=doc_context,
        system_information=location_data,
        doc_metadata=meta_context
    )
   
    return top_match, context_template, modified_query
 
@app.route("/preview/<id>/<filename>")
def preview_file(filename, id):
    UPLOAD_FOLDER = f"temp/{id}"
    app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, mimetype='application/pdf')
 
@app.route("/download/<id>/<filename>")
def download_file(filename, id):
    UPLOAD_FOLDER = f"temp/{id}"
    app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, as_attachment=True)
 
@app.route("/get_user_details", methods=["GET"])
def get_user_details():
    try:
        token = request.headers.get("Authorization")
        if not token:
            return jsonify({"error": "Missing token"}), 401
   
        user_id = verify_token(token)
        if not user_id:
            return jsonify({"error": "Invalid or expired token"}), 401
   
        username = get_username(user_id)
        if not username:
            return jsonify({"error": "User not found"}), 404
   
        return jsonify({"username": username.upper()})
    except Exception as e:
        print(f"Unexpected error in /get_user_details: {e}")
        return jsonify({"error": "An unexpected error occurred"}), 500
 
# API Routes
@app.route("/register", methods=["POST"])
def register():
    data = request.get_json()
    username = data.get("username")
    password = data.get("password")
    if not username:
        return jsonify({"error": "Username is required"}), 400
    if not password:
        return jsonify({"error": "Password is required"}), 400
 
    success, error = register_user(username, password)
    if not success:
        return jsonify({"error": error}), 400 if "already exists" in error else 500
   
    return jsonify({"message": "User registered successfully"}), 201
 
@app.route("/login", methods=["POST"])
def login():
    data = request.get_json()
    username = data.get("username")
    password = data.get("password")
 
    if not username:
        return jsonify({"error": "Username is required ❗"}), 400
    if not password:
        return jsonify({"error": "Password is required ❗"}), 400
 
    try:
        user = get_user_by_password(password)
        if not user:
            return jsonify({"error": "User not found"}), 404
 
        user_name, user_id = user
        if user_name != username:
            return jsonify({"error": "Invalid username (or) password"}), 404
       
        chat_history = get_chat_history(user_id)
        chat_history = [{"role": row["role"], "content": row["content"]} for row in chat_history.get(1, [])][-50:]  # Get most recent 50 messages from default chat
 
        token = generate_token(user_id)
        return jsonify({
            "message": "Login successful",
            "token": token,
            "chat_history": chat_history
        })
    except Exception as e:
        return jsonify({"error": f"Error during login: {e}"}), 500
   
@app.route("/upload_documents", methods=["POST"])
def upload_documents():
    token = request.headers.get("Authorization")
    if not token:
        return jsonify({"error": "Missing token"}), 401
 
    user_id = verify_token(token)
    if not user_id:
        return jsonify({"error": "Invalid or expired token"}), 401
 
    if "files" not in request.files:
        return jsonify({"error": "No files uploaded"}), 400
   
    uploaded_files = request.files.getlist("files")
 
    try:
        current_used_space = get_used_storage(user_id)
        existing_files = get_existing_files(user_id)
        updated_files = []
        total_files = 0
        file_hashes = set()
        temp_folder = os.path.join("temp", str(user_id))
        os.makedirs(temp_folder, exist_ok=True)
 
        for file in uploaded_files:
            file_name = file.filename
            file_contents = file.read()
 
            if file_name == "" or not allowed_file(file_name):
                return jsonify({
                    "error": f"❗🥲 Invalid file format: {file_name}🙆🏻‍♂️"
                }), 400
 
            file_extension = file_name.rsplit(".", 1)[1].lower()
            file_path = os.path.join(temp_folder, file_name)
 
            if file_name in existing_files:
                os.remove(file_path)
                delete_file(user_id, file_name)
                updated_files.append(file_name)
 
            with open(file_path, "wb") as f:
                f.write(file_contents)
 
            total_files += 1
            file_size = os.path.getsize(file_path)
            file_hash = compute_file_hash(file_path)
 
            if file_hash in file_hashes:
                os.remove(file_path)
                return jsonify({
                    "error": f"❗🥲 Duplicate file detected: {file_name}. Remove duplicates from the list."
                }), 400
 
            if current_used_space + file_size > MAX_USER_STORAGE:
                os.remove(file_path)
                return jsonify({
                    "error": "❗🥲 Storage limit exceeded🗄️ (30MB)"
                }), 400
 
            file_hashes.add(file_hash)
 
        for file in uploaded_files:
            file_name = file.filename
            file_contents = file.read()
            file_extension = file_name.rsplit(".", 1)[1].lower()
            file_path = os.path.join(temp_folder, file_name)
            file_size = os.path.getsize(file_path)
 
            extracted_text = extract_text(file_path, file_extension)
            chat_response = chat_model.invoke(f"Text format of single resume doc:{extracted_text},##Format this resume without bullet points and **,##Don't give any extra text or ** or any highlighting string in response").content
            doc = Document(page_content=chat_response)
 
            if file_name.lower().endswith(".pdf"):
                preview_link = f"http://localhost:5000/preview/{user_id}/{file_name}"
            elif file_name.lower().endswith(".docx"):
                preview_link = f"http://localhost:5000/download/{user_id}/{file_name}"
 
            store_file_metadata(user_id, file_name, file_extension, file_size, preview_link)
            store_document_embedding(user_id, file_name, doc.page_content, embeddings_model.embed_query(doc.page_content))
 
        if total_files == 1:
            info_msg = "1 file uploaded successfully👍🏻"
        else:
            info_msg = "All files uploaded successfully👍🏻"
 
        if len(updated_files) == total_files:
            if len(updated_files) == 1:
                return jsonify({"message":"One file updated successfully 😉","files_updated":updated_files}),200
            else:
                return jsonify({"message":f"{len(updated_files)} files updated successfully 😉","files_updated":updated_files}),200
        elif len(updated_files) == 1:
            return jsonify({"message":f"{info_msg} and {len(updated_files)} file updated 😉","files_updated":updated_files}),200
        elif len(updated_files) != 0:
            return jsonify({"message":f"{info_msg} and {len(updated_files)} files updated 😉","files_updated":updated_files}),200
        else:
           return jsonify({"message":info_msg}),200
    except Exception as e:
        return jsonify({"error": f"Upload failed: {str(e)}"}), 500
    
@app.route("/ask", methods=["POST"])
def ask():
    try:
        token = request.headers.get("Authorization")
        if not token:
            return jsonify({"error": "Missing token"}), 401
   
        user_id = verify_token(token)
        if not user_id:
            return jsonify({"error": "Invalid or expired token"}), 401
   
        data = request.get_json()
        question = data.get("question")
        newchat_id = data.get("chatid")
        if not question:
            return jsonify({"error": "Question is required"}), 400
       
        relevant_docs, context_template, mdfy_question = retrieve_similar_context(user_id, newchat_id, question)
        chat_response = chat_model.invoke(f"{context_template}\nRelevant Docs: {relevant_docs}\n\nUser Query: {mdfy_question}").content
       
        # Store both user question and AI response
        store_chat_embedding(user_id, newchat_id, question, "user", embeddings_model.embed_query(question))
        store_chat_embedding(user_id, newchat_id, chat_response, "assistant", embeddings_model.embed_query(chat_response))
       
        return jsonify({"answer": chat_response})
    except Exception as e:
        print(f"Error in /ask route: {e}")
        return jsonify({"error": "An error occurred while processing your request"}), 500
 
@app.route("/get_chat_history", methods=["GET"])
def get_chat_history_route():
    try:
        token = request.headers.get("Authorization")
        if not token:
            return jsonify({"error": "Missing token"}), 401
 
        user_id = verify_token(token)
        if not user_id:
            return jsonify({"error": "Invalid or expired token"}), 401
 
        chat_data = get_chat_history(user_id)
        return jsonify({"chat_history": chat_data}), 200
    except Exception as e:
        print(f"Error in /get_chat_history route: {e}")
        return jsonify({"error": "An error occurred while fetching chat history"}), 500
 
@app.route("/user_files", methods=["GET"])
def get_user_files_route():
    try:
        token = request.headers.get("Authorization")
        if not token:
            return jsonify({"error": "Missing token"}), 401
   
        user_id = verify_token(token)
        if not user_id:
            return jsonify({"error": "Invalid or expired token"}), 401
   
        files = get_user_files(user_id)
        return jsonify({"files": files})
    except Exception as e:
        print(f"Error in /user_files route: {e}")
        return jsonify({"error": "An error occurred while fetching user files"}), 500
 
if __name__ == "__main__":
    # socketio.run(app,allow_unsafe_werkzeug=True)
    app.run(debug=True)